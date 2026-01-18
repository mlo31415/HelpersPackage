import docx
from docx.shared import Inches
from docx import text
from docx.text import paragraph
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.document import Document as DocumentObject


# Add a new style (or replace an existing style) in a document.
def AddStyleToDocument(document: DocumentObject, styleName: str, faceName: str, size: Pt=Pt(12), bold: bool=False, italic: bool=False, color: RGBColor=RGBColor(255, 255, 255), alignment=WD_ALIGN_PARAGRAPH.LEFT):

    styles = document.styles

    # 1. Add a new paragraph style named 'CustomStyle'
    custom_style = styles.add_style(styleName, WD_STYLE_TYPE.PARAGRAPH)
    custom_style.base_style = styles['Normal'] # Inherit from 'Normal'

    # 2. Define character formatting
    font = custom_style.font
    font.name = faceName
    font.size = size
    font.bold = bold
    font.italic = italic
    font.color.rgb = color

    # Define paragraph formatting (center alignment)
    paragraph_format = custom_style.paragraph_format
    paragraph_format.alignment = alignment


# Create a docx and a .txt version for the pocket program
def AppendStyledParaToDoc(doc: DocumentObject, txt: str, style="Normal"):
    para=doc.add_paragraph(txt, style=style)


def AppendTextToPara(para: docx.text.paragraph.Paragraph, txt: str):
    para.add_run(txt)


def AppendStyledTextToPara(para: docx.text.paragraph.Paragraph, txt: str, charstyle: str|None=None, bold: bool|None = None, italic: bool|None=None, size: float|None = None,
    indent: float|None = None, font: str|None = None):
    run=para.add_run(txt)
    if charstyle is not None:
        run.style = charstyle

    if bold is not None:
        run.bold=bold
    if italic is not None:
        run.italic=italic
    if font is not None:
        run.font.name=font
    if size is not None:
        run.font.size=Pt(size)
    if indent is not None:
        para.paragraph_format.left_indent=Inches(indent)

    para.paragraph_format.line_spacing=1
    para.paragraph_format.space_after=0